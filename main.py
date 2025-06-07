import asyncio
from config import FILE_PATH
from loader import load_pdf_pages
from splitter import split_documents
from vector_store import create_vector_store
from search import search
from rerank import create_bm25, search_with_rerank
from qa import build_qa_pipeline
from docx import Document
import time

async def main():
    pages = await load_pdf_pages(FILE_PATH)
    split_docs = split_documents(pages)

    vector_store, index, embeddings = create_vector_store(split_docs)
    bm25 = create_bm25(split_docs)

    print(search("What is Agentic AI?", embeddings, index, vector_store))
    print(search_with_rerank("What is Agentic AI?", bm25, split_docs))

    rag = build_qa_pipeline(vector_store)
    response = rag.invoke("What is foundational understanding of AI Agent")

    doc = Document()
    doc.add_heading('RAG Output', level=1)
    doc.add_paragraph("Question: What is foundational understanding of AI Agent")
    doc.add_paragraph("Answer:")
    doc.add_paragraph(response['result'])
    doc.save("rag_output.docx")
    print("Output saved to rag_output.docx")

if __name__ == "__main__":
    start_time = time.time()
    asyncio.run(main())
    end_time = time.time()
    print(f"Execution time: {end_time - start_time:.2f} seconds")